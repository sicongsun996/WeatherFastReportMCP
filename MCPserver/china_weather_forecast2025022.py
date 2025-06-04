from mcp.server.fastmcp import FastMCP
import matplotlib.pyplot as plt
import pandas as pd 
import geopandas as gpd
import matplotlib.pyplot as plt
import geojson
from adjustText import adjust_text
import requests
import pandas as pd 

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement




# 和风天气API常量,此处需要输入你的和风天气API BASE 还有 API KEY
HEWEATHER_API_BASE = "https://YOUR-API-BASE.re.qweatherapi.com"
params = {'key':'YOUR-API-KEY'}
headers = {"Authorization": "Bearer your_token"}
# 初始化FastMCP服务器
mcp = FastMCP("weather_fast_report")
#matplotlib中文显示设置
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']  # 添加中文字体名称
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

#中国城市adcode对照表
df_adcode  = pd.read_csv('files/China-City-List-latest.csv', header=1)

def get_api_data(url, params=None, headers=None):
    try:
        response = requests.get(url, params=params, headers=headers)
        response.raise_for_status()  # 检查HTTP错误
        return response.json()  # 返回JSON数据
    except requests.exceptions.RequestException as e:
        print(f"请求错误: {e}")
        return None
def get_location_id(location:str):
    #获取location_id
    url = f"{HEWEATHER_API_BASE}/geo/v2/city/lookup?location={location}"
    data = get_api_data(url,params=params)#获取地区信息
    if not data or data.get("code") != "200":
        raise Exception("获取位置信息失败，请检查地区名称或稍后重试!") 
    location_id = data['location'][0]['id']
    return location_id
def get_areas(location_id):
    AD_code = df_adcode.loc[df_adcode['Location_ID']==location_id,'AD_code'].values[0]
    with open(f'files/geo-json/{AD_code}.txt',encoding='utf-8') as f:
        geo_json = geojson.load(f)
    adcode_list = []
    for i in range(len(geo_json['features'])):
        geo_json['features'][i]['id'] = geo_json['features'][i]['properties']['adcode']
        adcode_list.append(geo_json['features'][i]['properties']['adcode'])
    
    locations = []
    for adcode in adcode_list:
        url = f"{HEWEATHER_API_BASE}/geo/v2/city/lookup?location={adcode}"
        data = get_api_data(url, params)
        if not data or data.get("code") != "200":
            raise Exception("获取位置信息失败，请检查地区名称或稍后重试")
        locations.append(pd.DataFrame(data['location']))
        location_df = pd.concat(locations,axis=0)
    location_df['AD_code'] = adcode_list
    show_geo = gpd.GeoDataFrame(geo_json['features'])
    cities = gpd.GeoDataFrame(show_geo.properties.to_list())
    show_geo['AD_code'] = cities['adcode']
    location_df = pd.merge(show_geo,location_df[['id','lat','lon','name','AD_code']],on='AD_code',how = 'inner')
    return location_df
def get_precip(location_id: str,time='24h',interval=None) -> str:
    """获取指定地点的降水预报信息
   
    Args:
        location: 需要查询地区的名称，支持文字、以英文逗号分隔的经度,纬度坐标（十进制，最多支持小数点后两位）、LocationID或Adcode（仅限中国城市）。
        time:需要查询多久后的天气预报，可选值：
            3d:3天预报
            7d:7天预报
            10d :10天预报
            15d :15天预报
            30d:30天预报
            24h:24小时预报
            72h:72小时预报
            168h:168小时预报
        interval:时间间隔，为整数。配合time实现预报切片，更灵活地获得不同天气尺度的预报。如“time='7d'，interval=5”则为获得5天天气预报；又如"time='24h',interval=5",则获得5小时天气预报
    Returns:
        格式化后的实时天气信息或错误信息
    """
    #
    if time in ['3d','7d','10d','15d','30d','24h','72h','168h']:
        url = f"{HEWEATHER_API_BASE}/v7/weather/{time}?location={location_id}"#天气预报
        forecast_info = get_api_data(url,params=params)

        if 'd' in time:
            forecast_df = pd.DataFrame(forecast_info['daily'])
        elif 'h' in time:
            forecast_df = pd.DataFrame(forecast_info['hourly'])
    else:
        raise Exception("输入时间有误") 
    if interval!= None:
        try:
            interval = int(interval)
            forecast_df_interval =  forecast_df.iloc[:interval]
        except:
            raise Exception('interval输入异常，应为int型')
    else:
        forecast_df_interval=forecast_df

    precip_min = forecast_df_interval['precip'].astype('float32').min()   
    precip_max = forecast_df_interval['precip'] .astype('float32').max() 
    precip_mean = forecast_df_interval['precip'] .astype('float32').mean()     
    return [precip_min,precip_max,precip_mean]
def add_font(element, font_name):
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    element._element.append(rPr)
 
@mcp.tool()
async def fastreport_in_word(location: str,time='24h',interval=None)->str:#后续封装为tool

    """获取指定地点的降水预报信息

Args:
    location: 需要查询地区的名称，支持文字、以英文逗号分隔的经度,纬度坐标（十进制，最多支持小数点后两位）、LocationID或Adcode（仅限中国城市）。
    time:需要查询多久后的天气预报，可选值：
        3d:3天预报
        7d:7天预报
        10d :10天预报
        15d :15天预报
        30d:30天预报
        24h:24小时预报
        72h:72小时预报
        168h:168小时预报
    interval:时间间隔，为整数。配合time实现预报切片，更灵活地获得不同天气尺度的预报。如“time='7d'，interval=5”则为获得5天天气预报；又如"time='24h',interval=5",则获得5小时天气预报
Returns:
    格式化后的实时天气信息或错误信息
    """
    location_id=get_location_id(location)
    precip_min,precip_max,precip_mean = get_precip(location_id)
    areas= get_areas(location_id)
    areas[['precip_min','precip_max','precip_mean']] = areas['id_y'].apply(lambda x : get_precip(x)).tolist()
    if interval == None:
        fast_report =f"""{location}地区{time}降水快报：全市平均降水{precip_mean:.2f}毫米,过程降水介于{precip_min:.2f}毫米~{precip_max:.2f}毫米之间；各区（县）降雨量（单位：毫米）:"""
    else:
        fast_report =f"""{location}地区{interval}{time[-1]}降水快报：全市平均降水{precip_mean:.2f}毫米,过程降水介于{precip_min:.2f}毫米~{precip_max:.2f}毫米之间；各区（县）最大降雨量（单位：毫米）："""
    count = 0
    for i in range(len(areas)):
        fast_report += f"{areas.iloc[i]['name']} {areas.iloc[i]['precip_max']:.2f};"
        count+=1
    fast_report=fast_report[:-1] +'。'
    fast_report=fast_report.replace("h","小时")
    fast_report=fast_report.replace("d","天")

    areas[['lon','lat']]= areas[['lon','lat']].astype('float32')

    areas = gpd.GeoDataFrame(areas)
    ax = areas.plot(edgecolor='black',figsize=(15, 10),column ='precip_max' ,cmap='Blues',legend=True, legend_kwds={"shrink":.5},)
    texts = [plt.text(x_, y_, text, fontsize=14,color= 'black',bbox=dict(boxstyle='round', ec='none', fc='white', pad=0.1,alpha = 0.5)) for x_, y_, text in zip(areas['lon'],areas['lat'],areas['name'])]
    adjust_text(texts)
    # plt.show()
    colorbar_ax = ax.get_figure().axes[-1] #to get the last axis of the figure, it's the colorbar axes
    colorbar_ax.set_title(f"单位：mm", size=10)
    plt.axis('off')
    pic_path = f'{location}预测图片.png'
    plt.savefig(pic_path,bbox_inches='tight')

    # 创建一个Word文档对象
    doc = Document()

    # 添加标题
    heading = doc.add_heading(level=1)
    run = heading.add_run(f"{location}天气快报")
    add_font(run, '黑体')  
    run.font.size = Pt(20)  # 设置字体大小

    para = doc.add_paragraph()
    run = para.add_run(fast_report)
    add_font(run, '仿宋')  
    run.font.size = Pt(16)  # 设置字体大小
    
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.first_line_indent=Pt(30)
    para = doc.add_paragraph()
    para.add_run().add_picture(pic_path, width=Inches(7))
    # 添加表格
    row = len(areas['name'])
    table = doc.add_table(rows=row+1, cols=4)  # 创建一个3x3的表格
    areas_table = areas[['name','precip_max','precip_min','precip_mean']]
    areas_table[['precip_max','precip_min','precip_mean']] = areas_table[['precip_max','precip_min','precip_mean']].astype('float32').round(2)
    # 填充表格内容

    for j in range(4):
        table.cell(0, j).text =['地区','最大降水量','最小降水量','平均降水量'][j]
    for i in range(row):
        for j in range(4):
            table.cell(i+1, j).text =str(areas_table.iloc[i,j])

            # 保存文档
    doc_path = f'{location}天气快报.docx'
    doc.save(doc_path)
    return doc_path


if __name__ == "__main__":
    # 启动MCP服务器
    mcp.run(transport='stdio')
